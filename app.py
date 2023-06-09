import streamlit as st
import openai
import json
from docx import Document
from docx.shared import Inches

with open('config.json') as f:
    config = json.load(f)
    
# Variables
openai.api_type = config["openai"]["api_type"]
openai.api_version = config["openai"]["api_version"]
openai.api_base = config["openai"]["api_base"]
openai.api_key = config["openai"]["api_key"]
engine = config["engine"]["engine"]


# pip install openai
# openai.api_key = "OPENAI_API_KEY"

# Define functions
def get_cv():
    cv_text = st.text_area('Paste Part of CV here')
    return cv_text

def get_job_description():
    job_description = st.text_area('Paste a part of Job Description here')
    return job_description

def generate_docx(text):
    document = Document()
    document.add_heading('Cover Letter', 0)
    document.add_paragraph(text)
    document.save('cover_letter.docx')
    
    
def cover_letter_generator(cv_excerpt, job_description_excerpt):
    prompt = f"""
                Generate a compelling cover letter using the details from my CV:
                {cv_excerpt} and the provided job description:
                {job_description_excerpt}. The cover letter should effectively and accurately  highlight my skills and experiences in relation to the job requirements.
                """
    
    completions = openai.Completion.create(
        engine=engine,
        prompt=prompt,
        max_tokens=512,
        n=1,
        stop=None,
        temperature=0.2,
    )
    response = completions.choices[0].text
    
    return response

def main():
    st.set_page_config(page_title="Project Kurusha", page_icon=":rocket:")
    st.title(" Cover Letter Generator")
    cv_text = get_cv()
    if cv_text:
        job_desc = get_job_description()
        if job_desc:
            if st.button('Generate Cover Letter'):
                response = cover_letter_generator(cv_text, job_desc)
                generate_docx(response)
                with open('cover_letter.docx', 'rb') as docx_file:
                    st.download_button('Download', data=docx_file, file_name='cover_letter.docx')
                

if __name__  == '__main__':
    main()      

    
    